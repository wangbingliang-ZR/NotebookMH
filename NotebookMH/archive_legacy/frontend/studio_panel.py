"""
frontend/studio_panel.py - Studio 面板 (右侧)

完全参照 Google NotebookLM Studio Panel 设计:
  - 顶部: Studio 标题
  - 功能卡片区: 音频概述 / 学习指南 / 简报文档 / 时间线 / FAQ / 闪卡
  - 笔记区: 用户保存的 AI 回复
"""

import logging
import asyncio
from typing import Any, Dict, List

try:
    import streamlit as st
except ImportError:  # pragma: no cover
    st = None  # type: ignore[assignment]

from utils.state_manager import binder
from utils.db_manager import db_pool

logger = logging.getLogger(__name__)

_SSK_NOTES = "nb_mh_studio_notes"

# ---------------------------------------------------------------------------
# NotebookLM 风格功能卡片定义
# ---------------------------------------------------------------------------

_STUDIO_TOOLS = [
    {
        "key": "summary",
        "title": "文档摘要",
        "desc": "自动提炼文档核心内容",
        "prompt_query": "请总结全部文档的核心内容",
        "prompt_gen": "请基于以下内容生成一份简洁的中文摘要（300字以内）：\n\n{context}",
        "system": "你是一个专业的学习助手，擅长总结文档。",
    },
    {
        "key": "faq",
        "title": "常见问题 FAQ",
        "desc": "从文档中提炼常见问题及答案",
        "prompt_query": "文档中的关键知识点和常见问题",
        "prompt_gen": "请基于以下内容生成5个常见问题及回答（FAQ格式，中文）：\n\n{context}",
        "system": "你是一个专业的学习助手，擅长提炼FAQ。",
    },
    {
        "key": "study_guide",
        "title": "学习指南",
        "desc": "生成核心概念清单与学习路径",
        "prompt_query": "核心概念和学习路径",
        "prompt_gen": "请基于以下内容生成一份学习指南，包括：\n1. 核心概念清单\n2. 学习路径建议\n3. 重点难点提示\n\n{context}",
        "system": "你是一个专业的学习规划师。",
    },
    {
        "key": "timeline",
        "title": "时间线",
        "desc": "按时间顺序梳理关键事件",
        "prompt_query": "按时间顺序列出关键事件和里程碑",
        "prompt_gen": "请基于以下内容，按时间顺序生成一份事件时间线（中文）：\n\n{context}",
        "system": "你是一个专业的信息整理助手，擅长按时间顺序梳理事件。",
    },
    {
        "key": "briefing",
        "title": "简报文档",
        "desc": "生成正式的简报摘要",
        "prompt_query": "关键要点和结论",
        "prompt_gen": "请基于以下内容生成一份简报文档，包含：背景、要点、结论与建议（中文，500字以内）：\n\n{context}",
        "system": "你是一个专业的文秘助手，擅长撰写简报。",
    },
    {
        "key": "flashcard",
        "title": "闪卡 / 记忆卡",
        "desc": "生成 Q&A 记忆卡片",
        "prompt_query": "需要记忆的关键知识点",
        "prompt_gen": '请基于以下内容生成10张闪卡。返回JSON数组：[{"question":"问题","answer":"答案"}]\n只返回JSON，不要其他文字。\n\n内容：\n{context}',
        "system": "你是教育助手。只返回JSON数组格式的闪卡。",
    },
    {
        "key": "mindmap",
        "title": "思维导图",
        "desc": "可视化概念关系图",
        "prompt_query": "核心概念及其关系",
        "prompt_gen": (
            "请基于以下内容生成 Mermaid 格式的思维导图代码。要求：\n"
            "1. 使用 mindmap 语法\n"
            "2. 层级不超过3层\n"
            "3. 只返回代码，不要 ```mermaid 标记\n"
            "格式：\n"
            "mindmap\n"
            "  root((主题))\n"
            "    概念A\n"
            "      子概念\n\n"
            "内容：\n{context}"
        ),
        "system": "你是思维导图助手。只返回 Mermaid mindmap 代码，不要任何其他文字或 markdown 标记。",
    },
    {
        "key": "datatable",
        "title": "数据表格",
        "desc": "从来源中提取结构化表格",
        "prompt_query": "关键数据和对比信息",
        "prompt_gen": (
            "请基于以下内容提取关键信息，整理为 Markdown 表格。\n"
            "要求：3-6列，5-15行，含表头。只返回表格。\n\n"
            "内容：\n{context}"
        ),
        "system": "你是数据整理助手。只返回 Markdown 格式表格。",
    },
    {
        "key": "quiz",
        "title": "测验",
        "desc": "生成选择题测验",
        "prompt_query": "关键知识点和概念",
        "prompt_gen": '请基于以下内容生成5道选择题。返回JSON数组：[{"question":"题目","options":["A.xx","B.xx","C.xx","D.xx"],"correct":"A","explanation":"解析"}]\n只返回JSON，不要其他文字。\n\n内容：\n{context}',
        "system": "你是测验生成助手。只返回JSON数组格式的选择题。",
    },
]


# ---------------------------------------------------------------------------
# 公共渲染接口
# ---------------------------------------------------------------------------

def render() -> None:
    """渲染右侧 Studio 面板（NotebookLM 风格）。"""
    if st is None:
        return

    # 初始化笔记列表
    if _SSK_NOTES not in st.session_state:
        st.session_state[_SSK_NOTES] = []

    # ── 标题区 ────────────────────────────────────────────
    st.markdown(
        '<p style="font-size:22px; font-weight:600; color:#202124; margin:0 0 4px 0;">Studio</p>',
        unsafe_allow_html=True,
    )
    st.caption("基于你的来源文件，一键生成各种学习材料")
    st.divider()

    # ── 功能卡片区 ────────────────────────────────────────
    vault_uuid = binder.get_state("vault_uuid", "")
    has_docs = False
    if vault_uuid:
        docs = db_pool.list_documents(vault_uuid)
        has_docs = bool(docs)

    if not has_docs:
        st.markdown(
            '<div style="background:#f8f9fa; border-radius:12px; padding:24px; text-align:center;">'
            '<p style="font-size:15px; color:#5f6368; margin:0;">上传来源文件后</p>'
            '<p style="font-size:15px; color:#5f6368; margin:4px 0 0;">即可使用以下工具生成学习材料</p>'
            '</div>',
            unsafe_allow_html=True,
        )
        st.write("")

    # 渲染功能卡片（2列网格）
    st.markdown(
        '<p style="font-size:15px; font-weight:600; color:#202124; margin:12px 0 8px;">生成工具</p>',
        unsafe_allow_html=True,
    )

    for i in range(0, len(_STUDIO_TOOLS), 2):
        cols = st.columns(2)
        for col_idx, col in enumerate(cols):
            tool_idx = i + col_idx
            if tool_idx >= len(_STUDIO_TOOLS):
                break
            tool = _STUDIO_TOOLS[tool_idx]
            with col:
                st.markdown(
                    f'<div style="background:#f8f9fa; border:1px solid #e8eaed; border-radius:10px; '
                    f'padding:14px; margin-bottom:8px; min-height:80px;">'
                    f'<p style="font-size:14px; font-weight:600; color:#202124; margin:0 0 4px;">{tool["title"]}</p>'
                    f'<p style="font-size:12px; color:#5f6368; margin:0;">{tool["desc"]}</p>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                btn_disabled = not has_docs
                if tool["key"] in ("flashcard", "quiz"):
                    with st.expander("设置参数", expanded=False):
                        count = st.slider("数量", 5, 30, 10, key=f"slider_{tool['key']}")
                        difficulty = st.selectbox("难度", ["基础", "中等", "困难"], key=f"diff_{tool['key']}")
                        theme = st.text_input("主题（可选）", key=f"theme_{tool['key']}")
                        if st.button("生成", key=f"btn_studio_{tool['key']}", use_container_width=True, disabled=btn_disabled):
                            custom_tool = dict(tool)
                            extra = f"请生成{count}道。难度：{difficulty}。"
                            if theme:
                                extra += f"主题：{theme}。"
                            custom_tool["prompt_gen"] = extra + "\n" + tool["prompt_gen"]
                            _generate_content(vault_uuid, custom_tool)
                else:
                    if st.button(
                        "生成",
                        key=f"btn_studio_{tool['key']}",
                        use_container_width=True,
                        disabled=btn_disabled,
                    ):
                        _generate_content(vault_uuid, tool)

    # ── 自定义生成 ────────────────────────────────────────
    st.divider()
    st.markdown(
        '<p style="font-size:15px; font-weight:600; color:#202124; margin:12px 0 8px;">自定义生成</p>',
        unsafe_allow_html=True,
    )
    custom_prompt = st.text_area(
        "描述你想生成的内容",
        key="studio_custom",
        height=80,
        placeholder="例如：写一篇博客 / 生成考前重点 / 对比分析...",
    )
    if custom_prompt and st.button("生成", key="btn_studio_custom"):
        custom_tool = {
            "key": "custom",
            "title": "自定义报告",
            "prompt_query": custom_prompt,
            "prompt_gen": f"用户要求：{custom_prompt}\n\n请基于以下来源内容完成用户的要求：\n\n{{context}}",
            "system": "你是一个全能助手，请按用户要求生成内容。用中文回答。",
        }
        _generate_content(vault_uuid, custom_tool)

    # ── 笔记区 ───────────────────────────────────────────
    st.divider()
    from utils.db_manager import db_pool
    notes = db_pool.list_notes(vault_uuid, binder.get_state("user_id", "anonymous")) if vault_uuid else []

    st.markdown(
        f'<p style="font-size:15px; font-weight:600; color:#202124; margin:0 0 8px;">笔记 ({len(notes)})</p>',
        unsafe_allow_html=True,
    )
    search_term = st.text_input("搜索笔记", key="note_search", placeholder="输入关键词...")
    if search_term:
        notes = [n for n in notes if search_term.lower() in n.title.lower() or search_term.lower() in n.content.lower()]

    if notes:
        md = "\n\n---\n\n".join(f"## {n.title}\n{n.content}" for n in notes)
        st.download_button("导出 Markdown", md.encode("utf-8"), "notes.md", "text/markdown")
        try:
            from docx import Document
            import io
            doc = Document()
            for n in notes:
                doc.add_heading(n.title, level=2)
                doc.add_paragraph(n.content)
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button("导出 Word", buf.getvalue(), "notes.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception:
            pass
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            for n in notes:
                pdf.add_page()
                pdf.set_font("Arial", "B", 14)
                safe_title = n.title.encode("latin-1", "replace").decode("latin-1")
                pdf.cell(0, 10, safe_title, ln=True)
                pdf.set_font("Arial", "", 12)
                for para in n.content.split("\n"):
                    safe_para = para.encode("latin-1", "replace").decode("latin-1")
                    pdf.multi_cell(0, 8, safe_para)
            pdf_bytes = pdf.output(dest="S").encode("latin-1")
            st.download_button("导出 PDF", pdf_bytes, "notes.pdf", "application/pdf")
        except Exception:
            pass

        for note in notes:
            with st.container():
                tag_html = f'<p style="font-size:12px; color:#5f6368; margin:6px 0 0;">标签: {note.tags}</p>' if note.tags else ''
                st.markdown(
                    f'<div style="background:#ffffff; border:1px solid #e8eaed; border-radius:10px; '
                    f'padding:14px; margin-bottom:8px;">'
                    f'<p style="font-size:13px; font-weight:600; color:#1a73e8; margin:0 0 6px;">{note.title}</p>'
                    f'<p style="font-size:14px; color:#202124; margin:0; line-height:1.5;">{note.content[:300]}</p>'
                    f'{tag_html}'
                    f'</div>',
                    unsafe_allow_html=True,
                )
                c1, c2, c3 = st.columns(3)
                with c1:
                    pinned = st.checkbox("置顶", value=bool(note.pinned), key=f"pin_note_{note.id}")
                    if pinned != bool(note.pinned):
                        db_pool.update_note(note.id, pinned=int(pinned))
                        st.rerun()
                with c2:
                    with st.expander("编辑", expanded=False):
                        new_title = st.text_input("标题", value=note.title, key=f"edit_title_{note.id}")
                        new_content = st.text_area("内容", value=note.content, key=f"edit_content_{note.id}", height=80)
                        new_tags = st.text_input("标签（逗号分隔）", value=note.tags or "", key=f"edit_tags_{note.id}")
                        if st.button("保存", key=f"btn_save_note_{note.id}"):
                            db_pool.update_note(note.id, title=new_title, content=new_content, tags=new_tags)
                            st.rerun()
                with c3:
                    if st.button("分享", key=f"btn_share_{note.id}"):
                        token = f"note-{note.id}"
                        st.code(f"https://notebooklm.example.com/share/{token}", language=None)
                    if st.button("删除", key=f"btn_del_note_{note.id}", help="删除笔记"):
                        db_pool.delete_note(note.id)
                        st.rerun()
    else:
        st.markdown(
            '<div style="background:#f8f9fa; border-radius:10px; padding:20px; text-align:center;">'
            '<p style="font-size:14px; color:#5f6368; margin:0;">暂无笔记</p>'
            '<p style="font-size:13px; color:#9aa0a6; margin:4px 0 0;">对话中的 AI 回复可保存到这里</p>'
            '</div>',
            unsafe_allow_html=True,
        )


# ---------------------------------------------------------------------------
# 公共保存接口
# ---------------------------------------------------------------------------

def save_note(title: str, content: str) -> None:
    """从对话面板保存一条笔记到 Studio。"""
    try:
        from utils.db_manager import db_pool
        from utils.state_manager import binder
        vault_uuid = binder.get_state("vault_uuid", "")
        user_id = binder.get_state("user_id", "anonymous")
        if vault_uuid:
            db_pool.save_note(vault_uuid, user_id, title, content)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# 内容生成引擎（统一入口）
# ---------------------------------------------------------------------------

def _generate_content(vault_uuid: str, tool: Dict[str, str]) -> None:
    """通用内容生成：检索 + LLM 生成 + 保存笔记。"""
    try:
        from core.rag_pipeline import get_pipeline
        from core.llm_engine import get_llm_engine

        pipeline = get_pipeline()
        llm = get_llm_engine()

        async def _run():
            chunks = await pipeline.retrieve(tool["prompt_query"], vault_uuid, top_k=10)
            if not chunks:
                return "当前笔记库暂无可检索的内容。请先上传文件。"
            context = "\n\n".join([c.get("chunk_text", "") for c in chunks])
            prompt = tool["prompt_gen"].format(context=context)
            # ask_simple 返回纯文本字符串（chat 返回 AIResponse 对象）
            return await llm.ask_simple(prompt, system_prompt=tool["system"])

        with st.spinner(f"正在生成「{tool['title']}」..."):
            result = asyncio.run(_run())

        if tool["key"] == "mindmap":
            import streamlit.components.v1 as components
            components.html(
                '<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>'
                + f'<div class="mermaid" style="background:white; padding:20px; border-radius:12px;">{result}</div>'
                + '<script>mermaid.initialize({startOnLoad:true});</script>',
                height=500,
            )
            save_note(tool["title"], result)
        elif tool["key"] == "flashcard":
            import json as _json
            try:
                cards = _json.loads(result)
                if isinstance(cards, list):
                    from utils.db_manager import db_pool as _db
                    _db.save_flashcards(vault_uuid, cards)
                    save_note(tool["title"], f"已生成 {len(cards)} 张闪卡")
                    st.rerun()
            except _json.JSONDecodeError:
                pass
            save_note(tool["title"], result)
            st.rerun()
        elif tool["key"] == "quiz":
            import json as _json
            try:
                questions = _json.loads(result)
                if isinstance(questions, list):
                    from utils.db_manager import db_pool as _db
                    _db.save_quiz_questions(vault_uuid, questions)
                    save_note(tool["title"], f"已生成 {len(questions)} 道测验题")
                    st.rerun()
            except _json.JSONDecodeError:
                pass
            save_note(tool["title"], result)
            st.rerun()
        else:
            save_note(tool["title"], result)
            st.rerun()
    except Exception as e:
        logger.error("Studio generation [%s] failed: %s", tool["key"], e)
        st.error(f"生成失败: {e}")
